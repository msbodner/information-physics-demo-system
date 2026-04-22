#!/usr/bin/env python3
"""
Build InformationPhysics_CIM_TradeSecret.docx
Confidential Information Memorandum — InformationPhysics.ai, LLC
Aligned with the One-Way Transaction-Grade NDA (April 2026)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x34, 0x60)   # InformationPhysics sidebar blue
TEAL   = RGBColor(0x0E, 0x6B, 0x6B)   # subsection headings
RED    = RGBColor(0xC0, 0x00, 0x00)   # Trade Secret / alerts
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
LGREY  = RGBColor(0xF2, 0xF2, 0xF2)   # shaded table rows
MGREY  = RGBColor(0xD9, 0xD9, 0xD9)   # borders
GOLD   = RGBColor(0xC9, 0xA8, 0x2C)   # accent / stage labels

OUT = os.path.join(os.path.dirname(__file__), "InformationPhysics_CIM_TradeSecret.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex: str):
    """Fill a table cell with a solid background colour (hex string e.g. '0F3460')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top/bottom/left/right = (color_hex, size_eighths)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, (color, sz) in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=None, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=NAVY, size=None, font="Calibri"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = font
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    return p


def add_body(doc, text, indent=False, size=10.5, color=BLACK, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, text, size=size, color=color)
    return p


def add_bullet(doc, text, bold_prefix=None, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    if bold_prefix:
        add_run(p, bold_prefix + " ", bold=True, size=size)
    add_run(p, text, size=size)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br


def insert_page_break(doc):
    """Add a real page break paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    run = OxmlElement('w:r')
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run.append(br)
    p._p.append(run)
    return p


def add_watermark_header(doc, text="CONFIDENTIAL — TRADE SECRET"):
    """Add a red header line to every page acting as a watermark banner."""
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False
    # Clear default content
    for p in header.paragraphs:
        p.clear()
    if not header.paragraphs:
        hp = header.add_paragraph()
    else:
        hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RED
    run.font.name = "Calibri"
    # thin red border below header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C00000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(doc, page_label="InformationPhysics.ai, LLC — CIM | CONFIDENTIAL"):
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(page_label)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "Calibri"


def navy_table_header(table, headers, col_widths_inches):
    """Format first row of a table as a navy header row."""
    row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(row.cells, headers)):
        set_cell_bg(cell, '0F3460')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        cell.width = Inches(col_widths_inches[i])


def data_row(table, row_idx, values, shade=False, col_widths_inches=None):
    row = table.rows[row_idx]
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        if shade:
            set_cell_bg(cell, 'F2F2F2')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if isinstance(val, tuple):
            text, bold = val
        else:
            text, bold = val, False
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if col_widths_inches:
            cell.width = Inches(col_widths_inches[i])


# ── Build Document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup: US Letter, 1" margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(0.85)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_watermark_header(doc, "InformationPhysics.ai, LLC — CONFIDENTIAL INFORMATION MEMORANDUM | TRADE SECRET")
    add_footer(doc, "InformationPhysics.ai, LLC  ·  CIM April 2026  ·  CONFIDENTIAL — Subject to NDA")

    # ── COVER PAGE ────────────────────────────────────────────────────────────

    # Confidentiality banner
    banner = doc.add_table(rows=1, cols=1)
    banner.style = 'Table Grid'
    bc = banner.cell(0, 0)
    set_cell_bg(bc, 'C00000')
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after  = Pt(6)
    br = bp.add_run("CONFIDENTIAL — SUBJECT TO EXECUTED NON-DISCLOSURE AGREEMENT")
    br.bold = True
    br.font.color.rgb = WHITE
    br.font.size = Pt(10)
    br.font.name = "Calibri"

    doc.add_paragraph()

    # Logo / company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("InformationPhysics.ai, LLC")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Ann Arbor, Michigan  ·  informationphysics.ai")
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEAL
    r2.font.name = "Calibri"

    doc.add_paragraph()

    # CIM Title
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("CONFIDENTIAL INFORMATION MEMORANDUM")
    rt.bold = True
    rt.font.size = Pt(26)
    rt.font.color.rgb = NAVY
    rt.font.name = "Calibri"

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run("Proprietary AI/Data Platform  ·  IP Acquisition Opportunity")
    rs.font.size = Pt(13)
    rs.font.color.rgb = TEAL
    rs.font.name = "Calibri"

    doc.add_paragraph()

    # Date / version block
    meta = doc.add_table(rows=4, cols=2)
    meta.style = 'Table Grid'
    meta_data = [
        ("Date", "April 2026"),
        ("Document", "CIM Rev 1.0"),
        ("Classification", "Stage B — Post-NDA Disclosure"),
        ("Prepared by", "InformationPhysics.ai, LLC"),
    ]
    col_w = [1.8, 4.2]
    for i, (k, v) in enumerate(meta_data):
        row = meta.rows[i]
        set_cell_bg(row.cells[0], 'E8EDF3')
        pk = row.cells[0].paragraphs[0]
        pk.paragraph_format.space_before = Pt(3)
        pk.paragraph_format.space_after  = Pt(3)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.name = "Calibri"
        rk.font.color.rgb = NAVY
        row.cells[0].width = Inches(col_w[0])
        pv = row.cells[1].paragraphs[0]
        pv.paragraph_format.space_before = Pt(3)
        pv.paragraph_format.space_after  = Pt(3)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        rv.font.name = "Calibri"
        row.cells[1].width = Inches(col_w[1])

    doc.add_paragraph()
    doc.add_paragraph()

    # Trade Secret notice
    ts_box = doc.add_table(rows=1, cols=1)
    ts_box.style = 'Table Grid'
    tc2 = ts_box.cell(0, 0)
    set_cell_bg(tc2, 'FFF2CC')
    set_cell_border(tc2, top=('C9A82C', 8), bottom=('C9A82C', 8), left=('C9A82C', 8), right=('C9A82C', 8))
    tp = tc2.paragraphs[0]
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(6)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr1 = tp.add_run("⚠  TRADE SECRET — PROPRIETARY AND CONFIDENTIAL\n")
    tr1.bold = True
    tr1.font.size = Pt(11)
    tr1.font.color.rgb = RGBColor(0x7B, 0x5C, 0x00)
    tr1.font.name = "Calibri"
    tr2 = tp.add_run(
        "This document contains trade secrets and proprietary information protected under the "
        "Defend Trade Secrets Act, 18 U.S.C. § 1836 et seq., and applicable state law. "
        "Disclosure, reproduction, or use is strictly prohibited without the prior written "
        "consent of InformationPhysics.ai, LLC. Recipients acknowledge that this document "
        "is protected by an executed Non-Disclosure Agreement and may only be accessed by "
        "Named Reviewers as defined therein."
    )
    tr2.font.size = Pt(9)
    tr2.font.color.rgb = RGBColor(0x5A, 0x42, 0x00)
    tr2.font.name = "Calibri"

    # Page break after cover
    insert_page_break(doc)

    # ── SECTION 1: IMPORTANT NOTICE ──────────────────────────────────────────

    add_heading(doc, "IMPORTANT NOTICE — CONDITIONS OF ACCESS", level=1, color=RED, size=14)

    notice_text = (
        'This Confidential Information Memorandum ("CIM") has been prepared solely for the '
        "use of parties who have executed the InformationPhysics.ai, LLC One-Way Transaction-Grade "
        'Non-Disclosure Agreement ("NDA"). Receipt of this document constitutes acknowledgement '
        "of the binding confidentiality obligations therein."
    )
    add_body(doc, notice_text)

    add_heading(doc, "Staged Disclosure Framework", level=2, color=NAVY)
    add_body(doc, "This CIM constitutes a Stage B disclosure under the NDA's staged disclosure protocol:")

    stage_tbl = doc.add_table(rows=5, cols=3)
    stage_tbl.style = 'Table Grid'
    navy_table_header(stage_tbl, ["Stage", "Trigger", "Content Included"], [0.8, 2.2, 3.0])
    stages = [
        ("A", "Public / Pre-NDA", "Product overview, published papers, general capabilities"),
        ("B", "Executed NDA", "This CIM — architecture, IP summary, commercial opportunity"),
        ("C", "Executed LOI", "Full technical specifications, source code access (VDR), financial detail"),
        ("D", "Closing", "Complete IP transfer package, credentials, migration support"),
    ]
    for i, (stg, trig, cont) in enumerate(stages, 1):
        row = stage_tbl.rows[i]
        shade = (i % 2 == 0)
        if shade:
            set_cell_bg(row.cells[0], 'F2F2F2')
            set_cell_bg(row.cells[1], 'F2F2F2')
            set_cell_bg(row.cells[2], 'F2F2F2')
        for j, txt in enumerate([stg, trig, cont]):
            p = row.cells[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            if j == 0:
                r.bold = True
                r.font.color.rgb = GOLD

    doc.add_paragraph()
    add_heading(doc, "AI / Machine-Learning Prohibition", level=2, color=RED)
    add_body(doc,
        "Pursuant to the NDA, recipients are strictly prohibited from using any portion of "
        "this document or associated disclosures to train, fine-tune, embed, or otherwise "
        "incorporate into any artificial intelligence or machine-learning system, including "
        "but not limited to large language models, vector databases, retrieval-augmented "
        "generation pipelines, or automated summarization tools."
    )
    add_heading(doc, "Virtual Data Room", level=2, color=NAVY)
    add_body(doc,
        "Stage C and D materials are made available exclusively through an access-controlled "
        "Virtual Data Room (VDR) with full audit logging, watermarked documents, MFA authentication, "
        "and named-reviewer credentialing as specified in the NDA. No materials beyond this CIM "
        "are available outside the VDR."
    )

    insert_page_break(doc)

    # ── SECTION 2: EXECUTIVE SUMMARY ─────────────────────────────────────────

    add_heading(doc, "1.  Executive Summary", level=1)

    add_body(doc,
        "InformationPhysics.ai, LLC is offering for acquisition the intellectual property, "
        "software platform, and associated know-how comprising the Information Physics Standard "
        "Model — a novel, physics-inspired architecture for artificial intelligence data management "
        "that resolves fundamental limitations in how modern AI systems store, relate, retrieve, "
        "and audit information."
    )
    add_body(doc,
        "The platform is production-deployed, fully operational, and available for immediate "
        "demonstration. It has been developed entirely by founder Dr. Michael Simon Bodner and "
        "is owned free and clear by InformationPhysics.ai, LLC with no encumbrances, third-party "
        "licenses, or equity obligations."
    )

    add_heading(doc, "Transaction Highlights", level=2)

    highlights = doc.add_table(rows=6, cols=2)
    highlights.style = 'Table Grid'
    navy_table_header(highlights, ["Item", "Detail"], [2.0, 4.0])
    h_data = [
        ("Transaction Type", "Full IP acquisition (platform, patents-pending, trade secrets, know-how)"),
        ("Stage", "Seed-to-IP — no institutional investors, no cap table complexity"),
        ("Platform Status", "Production-deployed on Railway cloud; Electron desktop app (macOS/Windows/Linux)"),
        ("IP Ownership", "100% InformationPhysics.ai, LLC — no third-party claims"),
        ("Founder", "Dr. Michael Simon Bodner — see Section 3"),
    ]
    for i, (k, v) in enumerate(h_data, 1):
        shade = (i % 2 == 0)
        data_row(highlights, i, [(k, True), (v, False)], shade=shade, col_widths_inches=[2.0, 4.0])

    doc.add_paragraph()

    add_heading(doc, "Strategic Value Proposition", level=2)
    bullets = [
        ("Novel IP Category:", "The Information Physics Standard Model defines a new ontological layer beneath generative AI — not a product built on AI, but a substrate that makes AI auditable, persistent, and domain-sovereign."),
        ("No Comparable System:", "The AIO/HSL/MRO three-layer architecture has no direct commercial equivalent. It is protected as trade secret and subject to patent-pending filings."),
        ("Immediate Deployment:", "The full-stack platform (FastAPI backend, Next.js frontend, PostgreSQL persistence) is production-ready and cloud-deployed."),
        ("Cross-Industry Applicability:", "Healthcare, legal, financial services, defense, and enterprise intelligence all require exactly the auditability and lineage guarantees the platform provides."),
        ("Founder Continuity Available:", "Dr. Bodner is available for consulting, advisory, or integration support post-acquisition."),
    ]
    for bp, bt in bullets:
        add_bullet(doc, bt, bold_prefix=bp)

    insert_page_break(doc)

    # ── SECTION 3: FOUNDER ────────────────────────────────────────────────────

    add_heading(doc, "2.  Founder & Inventor", level=1)

    add_heading(doc, "Dr. Michael Simon Bodner", level=2, color=TEAL)

    add_body(doc,
        "Dr. Michael Simon Bodner (born April 4, 1948, Brooklyn, New York) is a physicist, "
        "technologist, and serial technology pioneer whose career spans six decades and every "
        "major computing revolution of the modern era. He is the sole inventor of the Information "
        "Physics Standard Model and the founder of InformationPhysics.ai, LLC, based in Ann Arbor, Michigan."
    )

    add_heading(doc, "Career Milestones", level=2)

    career_tbl = doc.add_table(rows=9, cols=3)
    career_tbl.style = 'Table Grid'
    navy_table_header(career_tbl, ["Period", "Organization / Role", "Contribution"], [1.0, 2.2, 2.8])
    career_data = [
        ("1968–1972", "NASA / Apollo Program", "Orbital mechanics & lunar lander guidance software; contributed calculations for Apollo 11"),
        ("1972–1978", "Raytheon Wayland Labs\nGrad: Boston University", "Phased-array radar development; M.S. Mathematical Physics"),
        ("1979–1981", "NASA (Visiting Scientist)", "LANDSAT data algorithms; Large Area Crop Inventory Experiment (LACIE)"),
        ("1980", "Traffic Central, Houston", "Designed real-time traffic information system — prototype of modern connected-data delivery"),
        ("1981–1985", "ComputerLand", "Advised stocking the Apple II; early LAN deployments for Bank of America"),
        ("1985–2000", "ComNet & enterprise CTO/CIO roles", "Early LAN design; legacy-to-networked transitions; cloud-era migrations"),
        ("2000–2020", "Multiple CTO/CIO roles", "Enterprise AI adoption, system architecture, data governance"),
        ("2022–present", "InformationPhysics.ai, LLC", "Founder & inventor — Information Physics Standard Model; author of three foundational research papers"),
    ]
    for i, row_data in enumerate(career_data, 1):
        shade = (i % 2 == 0)
        data_row(career_tbl, i, list(row_data), shade=shade, col_widths_inches=[1.0, 2.2, 2.8])

    doc.add_paragraph()
    add_body(doc,
        "Dr. Bodner is also the author of science-fiction novels exploring asteroid mining, "
        "artificial intelligence, and quantum entanglement — work that reflects the same rigorous "
        "physics foundation that underlies the Information Physics platform. His career represents "
        "an unbroken thread from NASA mission-critical computing to the frontier of AI architecture."
    )

    insert_page_break(doc)

    # ── SECTION 4: TECHNOLOGY ─────────────────────────────────────────────────

    add_heading(doc, "3.  Technology Overview", level=1)

    add_heading(doc, "The Core Paradigm: From Schema-First to Preserve-First", level=2)
    add_body(doc,
        "The foundational thesis of Information Physics is the transition from schema-first to "
        "preserve-first information architecture. Traditional AI systems require designers to "
        "decide which questions matter before storing data — an approach that is efficient for "
        "structured transactions but brittle in the face of generative AI's unpredictable, "
        "open-ended queries."
    )
    add_body(doc,
        "Information Physics inverts this: information is captured and preserved in its full, "
        "contextual state. Queries, analytics, and AI prompts are treated as measurements applied "
        "to this preserved state — enabling the system to answer questions that were never "
        "anticipated at ingestion time."
    )

    add_heading(doc, "The Three-Layer Standard Model", level=2)

    layers_tbl = doc.add_table(rows=4, cols=3)
    layers_tbl.style = 'Table Grid'
    navy_table_header(layers_tbl, ["Layer", "Name", "Function"], [1.0, 1.8, 3.2])
    layers_data = [
        ("Layer 1", "AIO — Associated Information Object", "Self-describing unit of information. 50-element bracket-notation schema. Carries semantic labels, provenance, temporal context, and identity. The quantum particle of the system."),
        ("Relational", "HSL — Hyper-Semantic Layer", "Precomputed pointer table linking AIOs via identity, similarity, temporal adjacency, and policy. 100-element schema. Transforms retrieval from search into traversal of a multidimensional semantic map."),
        ("Layer 2", "MRO — Memory Result Object", "Persisted retrieval episode. Records query, synthesized result, search terms (JSONB), and lineage. Enables recursive learning and auditable AI memory."),
    ]
    for i, row_data in enumerate(layers_data, 1):
        shade = (i % 2 == 0)
        data_row(layers_tbl, i, list(row_data), shade=shade, col_widths_inches=[1.0, 1.8, 3.2])

    doc.add_paragraph()

    add_heading(doc, "The Three Search Modes", level=2)
    add_bullet(doc, "All AIOs loaded as context; synthesizes a broad answer across the entire information universe. Optimal for discovery and exploratory queries.", bold_prefix="Broad Chat (ChatAIO):")
    add_bullet(doc, "Four-phase algebraic pipeline: (1) parse query → extract key terms; (2) match HSLs to identify semantic neighborhood; (3) gather linked AIOs; (4) synthesize a grounded, lineage-tracked answer. Optimal for precision retrieval.", bold_prefix="AIO Search:")
    add_bullet(doc, "Deterministic pipeline based on Paper III methodology. Exact semantic matching against the information substrate with full audit trail. Optimal for compliance and governed retrieval.", bold_prefix="Substrate Chat:")

    add_heading(doc, "The Seven Foundational Laws", level=2)
    laws = [
        ("1. Conservation of Provenance", "Every answer must trace back to its source. No derived result may exist without explicit lineage to the observations that produced it."),
        ("2. Contextual Meaning", "No measurement is valid without an explicit context frame. The same data may yield different answers under different measurement conditions."),
        ("3. Reversibility through Versioning", "Transformations create new versions rather than overwriting the original. The information universe only grows; it never silently loses state."),
        ("4. Entropy Accounting", "The system records what was lost or compressed during a measurement. Every summarization acknowledges its information cost."),
        ("5. Instrument Declaration", "Every answer must state which model or tool produced it. The measurement instrument is part of the measurement record."),
        ("6. Topological Boundedness", "Reasoning must define the boundaries of the evidence. Every answer specifies the neighborhood of the information universe it traversed."),
        ("7. Policy-Constrained Observability", "Access and visibility are governed by the physics of the system. Not all observers may measure all states; policy is a physical constraint."),
    ]
    for law, desc in laws:
        add_bullet(doc, desc, bold_prefix=law)

    add_heading(doc, "The Five-Layer Engineering Architecture", level=2)
    arch = [
        ("CORTEX (Cerebral Cortex)", "High-level reasoning and synthesis. Orchestrates LLM-powered analysis, multi-step inference, and answer generation."),
        ("HIPPOCAMPUS", "Object registry and multi-dimensional indexing. Stores AIOs, HSLs, and MROs with full provenance."),
        ("THALAMUS", "Routing layer — directs queries to the correct measurement instruments and synthesis engines."),
        ("BASAL GANGLIA", "Policy enforcement and action gating. Tenant isolation, access controls, admissibility rules."),
        ("CEREBELLUM", "Automated data transformation pipeline. Handles ingestion, AIO conversion, HSL precomputation."),
    ]
    for layer, desc in arch:
        add_bullet(doc, desc, bold_prefix=layer + ":")

    insert_page_break(doc)

    # ── SECTION 5: PLATFORM & FEATURES ───────────────────────────────────────

    add_heading(doc, "4.  Platform & Features", level=1)

    add_heading(doc, "Data Ingestion & Conversion", level=2)
    add_bullet(doc, "Drag-and-drop CSV or PDF upload converts raw data into AIO objects in real time.")
    add_bullet(doc, "PDF-to-CSV pipeline uses Claude AI to extract structured tabular data from unstructured documents.")
    add_bullet(doc, "AIO bracket notation ([Key.Value]) encodes semantic self-description at the object level.")
    add_bullet(doc, "Automated HSL builder precomputes the full semantic link graph from a batch of AIOs.")

    add_heading(doc, "ChatAIO — AI Retrieval Interface", level=2)
    add_bullet(doc, "Full-screen conversational interface with three independently selectable search modes.")
    add_bullet(doc, "MRO persistence: every retrieval episode is saved with query, result, and lineage for audit.")
    add_bullet(doc, "PDF export of any retrieval session for compliance and reporting.")
    add_bullet(doc, "Saved prompts library for repeatable, governed query patterns.")

    add_heading(doc, "System Management", level=2)
    add_bullet(doc, "10-tab admin panel: users, roles, AIO data, HSL data, API key, saved CSVs, saved AIOs, saved prompts, information elements, architecture.")
    add_bullet(doc, "Row-level security (PostgreSQL RLS) for full multi-tenant isolation by X-Tenant-Id header.")
    add_bullet(doc, "Role-based access control with named users and permission levels.")
    add_bullet(doc, "MRO ranking engine: score = Jaccard similarity × temporal decay × confidence weight.")

    add_heading(doc, "Deployment Targets", level=2)
    add_bullet(doc, "Cloud: Railway.app production deployment (auto-deploy from GitHub main branch).")
    add_bullet(doc, "Self-hosted: Docker Compose (PostgreSQL + FastAPI backend + Next.js frontend).")
    add_bullet(doc, "Desktop: Electron app (macOS DMG, Windows EXE, Linux AppImage) with bundled Python runtime, PostgreSQL, and standalone Next.js — zero external dependencies.")

    insert_page_break(doc)

    # ── SECTION 6: TECHNOLOGY STACK ──────────────────────────────────────────

    add_heading(doc, "5.  Technology Stack", level=1)

    stack_tbl = doc.add_table(rows=9, cols=3)
    stack_tbl.style = 'Table Grid'
    navy_table_header(stack_tbl, ["Component", "Technology", "Notes"], [1.5, 2.0, 2.5])
    stack_data = [
        ("Frontend", "Next.js 16 + React 19", "Turbopack dev; standalone production build"),
        ("UI Framework", "Tailwind CSS + shadcn/ui", "Radix UI primitives, Lucide icons"),
        ("Backend API", "FastAPI (Python 3.10+)", "~60KB main.py; 60+ endpoints"),
        ("Database", "PostgreSQL 15", "RLS, JSONB, 11 migration files"),
        ("AI / LLM", "Anthropic Claude API", "claude-3-5-sonnet; PDF extraction & search synthesis"),
        ("Desktop", "Electron", "Bundles Python 3.12, Node.js, PostgreSQL; cross-platform"),
        ("Cloud Deploy", "Railway.app", "Frontend + API + PostgreSQL services; auto-deploy from GitHub"),
        ("Containerization", "Docker Compose", "Self-hosted alternative; single-command startup"),
    ]
    for i, row_data in enumerate(stack_data, 1):
        shade = (i % 2 == 0)
        data_row(stack_tbl, i, list(row_data), shade=shade, col_widths_inches=[1.5, 2.0, 2.5])

    insert_page_break(doc)

    # ── SECTION 7: INTELLECTUAL PROPERTY ─────────────────────────────────────

    add_heading(doc, "6.  Intellectual Property Summary", level=1)

    add_body(doc,
        "The IP portfolio of InformationPhysics.ai, LLC comprises eight distinct and separable "
        "intellectual property assets. All are owned free and clear, with no third-party "
        "claims, encumbrances, open-source license obligations, or equity-linked IP provisions."
    )

    ip_tbl = doc.add_table(rows=9, cols=3)
    ip_tbl.style = 'Table Grid'
    navy_table_header(ip_tbl, ["ID", "IP Asset", "Description"], [0.6, 2.0, 3.4])
    ip_data = [
        ("IP-01", "Information Physics Standard Model", "The foundational theoretical framework — seven laws, three-layer data model, five-layer engineering architecture. Protected as trade secret; patent filings in preparation."),
        ("IP-02", "AIO Schema (50-Element Bracket Notation)", "The canonical self-describing observation object format. Novel information primitive not present in any known prior art."),
        ("IP-03", "HSL Schema (100-Element Pointer Table)", "Precomputed semantic link graph format enabling O(1) neighborhood traversal. Trade secret."),
        ("IP-04", "MRO Schema + Ranking Algorithm", "Episodic memory persistence format and Jaccard × decay × confidence scoring formula. Trade secret."),
        ("IP-05", "Four-Phase AIO Search Algebra", "Parse → HSL Match → AIO Gather → Synthesize pipeline. Novel retrieval algebra with no commercial equivalent."),
        ("IP-06", "Full-Stack Platform Source Code", "~200KB Python (FastAPI), TypeScript (Next.js/React), SQL migrations. Proprietary, no open-source obligations on core logic."),
        ("IP-07", "Three Foundational Research Papers", "Academic-grade documentation of the theoretical model, engineering architecture, and substrate retrieval methodology."),
        ("IP-08", "Brand & Domain", "InformationPhysics.ai domain, brand identity, and associated trademarks."),
    ]
    for i, row_data in enumerate(ip_data, 1):
        shade = (i % 2 == 0)
        data_row(ip_tbl, i, list(row_data), shade=shade, col_widths_inches=[0.6, 2.0, 3.4])

    doc.add_paragraph()
    add_body(doc,
        "Note: Patent applications are in preparation. The platform is currently protected "
        "exclusively under trade secret law (DTSA, 18 U.S.C. § 1836) and the executed NDA. "
        "Any disclosure, reverse engineering, or unauthorized use constitutes a federal trade "
        "secret misappropriation claim in addition to breach of contract."
    )

    insert_page_break(doc)

    # ── SECTION 8: MARKET OPPORTUNITY ────────────────────────────────────────

    add_heading(doc, "7.  Market Opportunity", level=1)

    add_heading(doc, "The Problem Information Physics Solves", level=2)
    problems = [
        ("AI Hallucination & Auditability:", "Current LLM systems cannot explain where an answer came from. Information Physics provides mandatory provenance and lineage at every retrieval step."),
        ("Data Sovereignty:", "Enterprise AI systems built on cloud-only vector stores cannot guarantee data isolation. The Information Physics architecture enables true multi-tenant substrate separation with RLS."),
        ("Memory Absence:", "LLMs have no persistent episodic memory across sessions. MROs solve this by persisting every retrieval episode as a governed, rankable memory object."),
        ("Schema Brittleness:", "Traditional databases require predefined schema. AIOs preserve information in full context, answering questions that were never anticipated at ingestion."),
        ("Compliance Gaps:", "AI-generated answers are currently unauditable in regulated industries. The Seven Laws and MRO lineage chain create a full audit trail from query to source."),
    ]
    for bp, bt in problems:
        add_bullet(doc, bt, bold_prefix=bp)

    add_heading(doc, "Target Acquirer Profiles", level=2)
    acquirers = [
        ("Enterprise AI Platforms", "Companies building vertical AI solutions (legal tech, healthtech, fintech) needing auditable, sovereign AI infrastructure."),
        ("Data Management Vendors", "Established database or data-warehouse vendors seeking a next-generation semantic layer above their storage products."),
        ("Defense & Intelligence Contractors", "Organizations requiring DTSA-compliant, air-gapped, fully auditable AI retrieval for classified or sensitive data environments."),
        ("Cloud Hyperscalers", "Platform teams seeking a novel AI data primitive to differentiate their AI/ML services portfolio."),
        ("Private Equity / Strategic Rollup", "PE firms building AI infrastructure portfolios who can scale the platform across multiple portfolio companies."),
    ]
    acq_tbl = doc.add_table(rows=6, cols=2)
    acq_tbl.style = 'Table Grid'
    navy_table_header(acq_tbl, ["Acquirer Type", "Strategic Fit"], [2.0, 4.0])
    for i, (aq, fit) in enumerate(acquirers, 1):
        shade = (i % 2 == 0)
        data_row(acq_tbl, i, [(aq, True), (fit, False)], shade=shade, col_widths_inches=[2.0, 4.0])

    insert_page_break(doc)

    # ── SECTION 9: DEAL STRUCTURE ─────────────────────────────────────────────

    add_heading(doc, "8.  Transaction Structure", level=1)

    add_heading(doc, "Available Transaction Structures", level=2)
    structures = [
        ("Full IP Acquisition:", "All eight IP assets transferred in a single transaction. Preferred structure for strategic buyers seeking complete platform ownership."),
        ("Licensed Deployment:", "Perpetual or term license to the platform and IP, with InformationPhysics.ai retaining ownership. Suitable for enterprises seeking deployment rights without acquisition."),
        ("Joint Venture / Co-Development:", "Collaborative arrangement for industry-specific customization and deployment, with shared IP development rights going forward."),
        ("Acqui-Hire:", "Acquisition of IP together with ongoing advisory or development engagement by Dr. Bodner."),
    ]
    for bp, bt in structures:
        add_bullet(doc, bt, bold_prefix=bp)

    add_heading(doc, "What Is Included in a Full Acquisition", level=2)
    inclusions = [
        "All source code (frontend, backend, database migrations, Electron packaging scripts)",
        "All three foundational research papers and associated documentation",
        "AIO, HSL, and MRO schema specifications",
        "Railway production deployment and domain transfer",
        "GitHub repository transfer",
        "Patent applications (in-process rights transferred)",
        "InformationPhysics.ai brand and domain",
        "Transition support period (negotiable) with Dr. Bodner",
    ]
    for item in inclusions:
        add_bullet(doc, item)

    add_heading(doc, "What Is Excluded", level=2)
    add_body(doc,
        "No employees, liabilities, accounts payable, or operating obligations transfer. "
        "This is a clean IP acquisition with no legacy encumbrances."
    )

    insert_page_break(doc)

    # ── SECTION 10: NEXT STEPS ────────────────────────────────────────────────

    add_heading(doc, "9.  Process & Next Steps", level=1)

    steps_tbl = doc.add_table(rows=5, cols=3)
    steps_tbl.style = 'Table Grid'
    navy_table_header(steps_tbl, ["Step", "Action", "NDA Stage"], [0.6, 4.0, 1.4])
    steps_data = [
        ("1", "Execute NDA → receive this CIM (Stage B) and schedule platform demonstration", "Stage A → B"),
        ("2", "Platform demonstration (live system); Q&A with Dr. Bodner; technical deep-dive available upon request", "Stage B"),
        ("3", "Execute Letter of Intent (LOI) → gain Stage C VDR access: full source code, financial detail, patent drafts", "Stage B → C"),
        ("4", "Definitive agreement negotiation and due diligence completion", "Stage C"),
    ]
    for i, row_data in enumerate(steps_data, 1):
        shade = (i % 2 == 0)
        data_row(steps_tbl, i, list(row_data), shade=shade, col_widths_inches=[0.6, 4.0, 1.4])

    doc.add_paragraph()

    add_heading(doc, "Contact", level=2)

    contact_tbl = doc.add_table(rows=4, cols=2)
    contact_tbl.style = 'Table Grid'
    contact_data = [
        ("Principal", "Dr. Michael Simon Bodner"),
        ("Organization", "InformationPhysics.ai, LLC  ·  Ann Arbor, Michigan"),
        ("Platform", "https://informationphysics.ai"),
        ("Engagement", "All inquiries subject to executed NDA. Contact via the platform or through counsel."),
    ]
    col_w2 = [1.5, 4.5]
    for i, (k, v) in enumerate(contact_data):
        row = contact_tbl.rows[i]
        set_cell_bg(row.cells[0], 'E8EDF3')
        pk = row.cells[0].paragraphs[0]
        pk.paragraph_format.space_before = Pt(3)
        pk.paragraph_format.space_after  = Pt(3)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.name = "Calibri"
        rk.font.color.rgb = NAVY
        row.cells[0].width = Inches(col_w2[0])
        pv = row.cells[1].paragraphs[0]
        pv.paragraph_format.space_before = Pt(3)
        pv.paragraph_format.space_after  = Pt(3)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        rv.font.name = "Calibri"
        row.cells[1].width = Inches(col_w2[1])

    insert_page_break(doc)

    # ── CLOSING CONFIDENTIALITY PAGE ─────────────────────────────────────────

    add_heading(doc, "CONFIDENTIALITY REMINDER", level=1, color=RED)

    closing_text = [
        "This Confidential Information Memorandum has been prepared exclusively for the addressee "
        "identified in the executed Non-Disclosure Agreement between InformationPhysics.ai, LLC "
        "and the receiving party. It is provided under the terms and conditions of that Agreement.",
        "",
        "By accessing this document, the recipient confirms that:",
        "",
        "1.  They are a Named Reviewer as defined in the NDA or are accessing this document under "
        "the express written authorization of a Named Reviewer.",
        "",
        "2.  They will not reproduce, distribute, disclose, or summarize any portion of this "
        "document to any person not listed as a Named Reviewer without prior written consent "
        "from InformationPhysics.ai, LLC.",
        "",
        "3.  They will not use any portion of this document or the information disclosed herein "
        "to train, fine-tune, embed, prompt-engineer, or otherwise incorporate into any artificial "
        "intelligence system, machine-learning model, vector database, or automated processing "
        "pipeline.",
        "",
        "4.  All physical and electronic copies of this document are subject to return or "
        "destruction upon request, with written certification of completion.",
        "",
        "Violations of these terms constitute breach of the NDA and misappropriation of trade "
        "secrets under the Defend Trade Secrets Act (18 U.S.C. § 1836 et seq.), entitling "
        "InformationPhysics.ai, LLC to seek injunctive relief, exemplary damages, and attorneys' "
        "fees without bond requirement.",
        "",
        "Whistleblower Immunity Notice (18 U.S.C. § 1833(b)): An individual shall not be held "
        "criminally or civilly liable under any Federal or State trade secret law for disclosure "
        "of a trade secret that is made in confidence to a Federal, State, or local government "
        "official or to an attorney solely for the purpose of reporting or investigating a "
        "suspected violation of law; or is made in a complaint or other document filed in a "
        "lawsuit or other proceeding, if such filing is made under seal.",
    ]

    for line in closing_text:
        if line == "":
            doc.add_paragraph()
        else:
            add_body(doc, line, size=10)

    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sig = p_sig.add_run("© 2026 InformationPhysics.ai, LLC. All rights reserved.\nTrade Secret — Protected under DTSA (18 U.S.C. § 1836)")
    r_sig.font.size = Pt(9)
    r_sig.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r_sig.font.name = "Calibri"
    r_sig.bold = True

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"✅  Saved: {OUT}")
    size_kb = os.path.getsize(OUT) // 1024
    print(f"   Size: {size_kb} KB")


if __name__ == "__main__":
    build()
