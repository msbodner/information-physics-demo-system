"""
Build InformationPhysics.ai Due Diligence Word Document
Trade Secret / Confidential
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree
import os

OUTPUT = os.path.join(os.path.dirname(__file__),
                      "InformationPhysics_DueDiligence_TradeSecret.docx")

# ── Color palette matching the paper style ─────────────────────────────────
NAVY       = RGBColor(0x0F, 0x34, 0x60)   # sidebar / header blue
TEAL       = RGBColor(0x0E, 0x6B, 0x6B)   # accent
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY   = RGBColor(0x66, 0x66, 0x66)
RED_TS     = RGBColor(0xCC, 0x00, 0x00)   # trade secret watermark / stamp

def set_cell_bg(cell, hex_color):
    """Set a table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_watermark(doc, text="TRADE SECRET"):
    """
    Add a diagonal text watermark to every section header.
    Uses WordprocessingML VML shape approach (works in Word + LibreOffice).
    """
    for section in doc.sections:
        header = section.header
        # Mark as linked-to-previous=False so it owns its own content
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run()

        # Build the VML picture element
        pic_xml = (
            '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office">'
            '<v:shape id="watermark" type="#_x0000_t136"'
            '  style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:300pt;'
            '         z-index:-251654144;mso-position-horizontal:center;'
            '         mso-position-horizontal-relative:margin;'
            '         mso-position-vertical:center;'
            '         mso-position-vertical-relative:margin;'
            '         rotation:315"'
            '  fillcolor="#CC0000" stroked="f">'
            '<v:fill on="t" type="solid"/>'
            '<v:textpath style="font-family:Arial;font-size:72pt;font-weight:bold;"'
            '  string="TRADE SECRET" trim="t"/>'
            '</v:shape>'
            '</w:pict>'
        )
        run._r.append(etree.fromstring(pic_xml))

def rule(doc, color_hex="0F3460"):
    """Add a coloured horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    rule(doc)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    return p

def body(doc, text, bold_phrases=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if not bold_phrases:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p

    # Split text around bold phrases
    remaining = text
    for phrase in bold_phrases:
        idx = remaining.find(phrase)
        if idx == -1:
            continue
        before = remaining[:idx]
        after  = remaining[idx + len(phrase):]
        if before:
            r = p.add_run(before)
            r.font.size = Pt(10.5); r.font.name = "Calibri"
            r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        r = p.add_run(phrase)
        r.bold = True
        r.font.size = Pt(10.5); r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        remaining = after
    if remaining:
        r = p.add_run(remaining)
        r.font.size = Pt(10.5); r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with navy header row."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        set_cell_bg(hdr_cells[i], "0F3460")
        r = hdr_cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10); r.font.name = "Calibri"
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = "F5F8FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = ""
            set_cell_bg(cells[ci], bg)
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5); r.font.name = "Calibri"

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height   = Inches(11)
section.page_width    = Inches(8.5)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Default font ────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── Watermark ───────────────────────────────────────────────────────────────
add_watermark(doc, "TRADE SECRET")

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Trade secret banner (top)
ts_p = doc.add_paragraph()
ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ts_p.paragraph_format.space_before = Pt(0)
ts_p.paragraph_format.space_after  = Pt(0)
pPr = ts_p._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"CC0000"); pPr.append(shd)
r = ts_p.add_run("⚠  TRADE SECRET — CONFIDENTIAL  ⚠")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()

# Logo / company block
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = logo_p.add_run("InformationPhysics.ai")
r.bold = True; r.font.size = Pt(28); r.font.name = "Calibri"
r.font.color.rgb = NAVY

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tagline_p.add_run("The Information Physics Standard Model")
r.font.size = Pt(13); r.font.name = "Calibri"
r.font.color.rgb = TEAL
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Thin rule
rule(doc)

# Document title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(24)
r = title_p.add_run("DUE DILIGENCE TECHNICAL REVIEW")
r.bold = True; r.font.size = Pt(22); r.font.name = "Calibri"
r.font.color.rgb = NAVY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle_p.add_run("AIO / HSL / MRO Demo System  ·  Version 3.5")
r.font.size = Pt(13); r.font.name = "Calibri"
r.font.color.rgb = MID_GREY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Metadata block
meta_lines = [
    ("Document Classification",  "TRADE SECRET — PROPRIETARY AND CONFIDENTIAL"),
    ("Prepared By",              "InformationPhysics.ai, LLC"),
    ("Author",                   "Michael Simon Bodner, Ph.D."),
    ("Version",                  "V3.5  |  April 2026"),
    ("Repository",               "github.com/msbodner/information-physics-demo-system"),
    ("Deployment",               "Railway Cloud  ·  Electron Desktop  ·  Docker"),
]
meta_tbl = doc.add_table(rows=len(meta_lines), cols=2)
meta_tbl.style = "Table Grid"
for i, (label, value) in enumerate(meta_lines):
    cells = meta_tbl.rows[i].cells
    bg = "E8EEF8" if i % 2 == 0 else "F5F8FF"
    set_cell_bg(cells[0], bg); set_cell_bg(cells[1], bg)
    cells[0].width = Inches(2.2); cells[1].width = Inches(3.6)
    rl = cells[0].paragraphs[0].add_run(label)
    rl.bold = True; rl.font.size = Pt(10); rl.font.name = "Calibri"; rl.font.color.rgb = NAVY
    rv = cells[1].paragraphs[0].add_run(value)
    rv.font.size = Pt(10); rv.font.name = "Calibri"
    if i == 0:
        rv.bold = True; rv.font.color.rgb = RED_TS

doc.add_paragraph()
rule(doc)
doc.add_paragraph()

# Confidentiality notice
notice_p = doc.add_paragraph()
notice_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = notice_p.add_run(
    "This document constitutes a trade secret of InformationPhysics.ai, LLC and contains "
    "proprietary, confidential information protected under applicable trade secret law, "
    "including the Defend Trade Secrets Act (18 U.S.C. § 1836 et seq.) and applicable "
    "state law. Unauthorized disclosure, reproduction, or use is strictly prohibited and "
    "may subject the violator to civil and criminal liability. Receipt of this document "
    "constitutes agreement to maintain its confidentiality."
)
r.font.size  = Pt(8.5)
r.font.name  = "Calibri"
r.font.color.rgb = MID_GREY
r.italic     = True

# Page break after cover
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

# Running header for body pages
section = doc.sections[0]
header  = section.header
hdr_p   = header.paragraphs[0]
hdr_p.clear()
hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r1 = hdr_p.add_run("InformationPhysics.ai — Due Diligence Review  |  ")
r1.font.size = Pt(8); r1.font.name = "Calibri"; r1.font.color.rgb = MID_GREY
r2 = hdr_p.add_run("TRADE SECRET")
r2.bold = True; r2.font.size = Pt(8); r2.font.name = "Calibri"; r2.font.color.rgb = RED_TS

h1(doc, "1.  Executive Summary")

body(doc,
    "The Information Physics Demo System is a proprietary, production-deployed AI platform "
    "built around a novel three-layer data model called the Information Physics Standard Model. "
    "The platform transforms structured data — CSV files and PDF documents — into self-describing "
    "semantic objects, links them through a precomputed relational index, and captures AI retrieval "
    "sessions as reusable knowledge objects. The result is a deterministic, bounded-cost information "
    "retrieval system that operates without vector embeddings or probabilistic similarity search.",
    bold_phrases=["Information Physics Standard Model", "deterministic, bounded-cost"])

body(doc,
    "The system embodies three original intellectual property constructs — the Associated "
    "Information Object (AIO), the Hyper-Semantic Layer (HSL), and the Memory Result Object (MRO) "
    "— implemented in a full production-grade stack with web, cloud, and desktop deployment targets. "
    "Each construct is formally documented in proprietary reference papers (Papers I, II, and III) "
    "that constitute protected trade secrets of InformationPhysics.ai, LLC.",
    bold_phrases=["Associated Information Object (AIO)", "Hyper-Semantic Layer (HSL)",
                  "Memory Result Object (MRO)"])

body(doc,
    "The platform is presently deployed on Railway (cloud production), GitHub (version-controlled "
    "CI/CD pipeline), and as a native Electron desktop application for macOS, Windows, and Linux. "
    "It uses Claude Sonnet 4.6 (Anthropic) as its AI reasoning engine and PostgreSQL 15 as its "
    "persistent data store. Current version: V3.5, April 2026.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PRODUCT PHILOSOPHY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "2.  Product Philosophy and Conceptual Foundation")

h2(doc, "2.1  The Information Physics Premise")

body(doc,
    "Conventional enterprise AI retrieval relies on floating-point vector embeddings and cosine "
    "similarity to surface semantically proximate content. This approach is probabilistic, opaque, "
    "computationally expensive at scale, and produces non-deterministic results that are difficult "
    "to audit or reproduce. As corpus size grows, embedding index costs grow proportionally.",
    bold_phrases=["probabilistic, opaque", "non-deterministic"])

body(doc,
    "Information Physics takes a fundamentally different position: structured information should be "
    "self-describing, explicitly linked, and retrievable through algebraic set operations. This "
    "philosophy is grounded in three proprietary reference papers (Papers I, II, III) that define "
    "the platform's standard model. The key claim is that a well-formed information structure, "
    "when properly indexed, allows a language model to reason from precomputed substrates — "
    "bounded, deterministic context bundles assembled without embedding search — achieving "
    "reliable, reproducible, cost-bounded inference.",
    bold_phrases=["self-describing, explicitly linked", "algebraic set operations",
                  "precomputed substrates"])

h2(doc, "2.2  The Standard Model — Three-Layer Architecture")

code_block(doc,
    "  Layer 1 (Atoms):      AIO  →  Self-describing [Key.Value] data objects\n"
    "  Relational Index:     HSL  →  Precomputed pointer tables linking AIOs\n"
    "  Layer 2 (Memory):     MRO  →  Episodic AI retrieval records reused as knowledge")

body(doc,
    "Each layer has a specific, non-overlapping role. AIOs are immutable, self-describing atomic "
    "records. HSLs are the connective tissue — precomputed maps that identify which AIOs share "
    "which element values, enabling O(1) neighborhood lookup at query time with no recomputation. "
    "MROs are institutional memory — when the AI answers a question, that complete episode "
    "(query, traversal path, matched evidence, synthesized answer) is captured as a first-class "
    "object and linked back into the HSL graph.",
    bold_phrases=["O(1) neighborhood lookup", "institutional memory"])

body(doc,
    "This three-layer approach is novel in the industry. No commercial system combines "
    "deterministic set-algebra retrieval with episodic memory linked back into a precomputed "
    "semantic index. The combination produces a self-improving knowledge graph: each answered "
    "query, when saved and linked, enriches future answers on related topics without retraining "
    "or re-embedding any model.",
    bold_phrases=["self-improving knowledge graph"])

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DATA MODEL
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "3.  Data Model — Technical Specification")

h2(doc, "3.1  Associated Information Objects (AIOs)")

body(doc,
    "AIOs are the atomic unit of the platform. Each AIO is a structured record in [Key.Value] "
    "bracket notation — a self-describing format that makes every field and its value unambiguous "
    "without requiring a separate schema lookup. AIOs are created automatically when CSV files "
    "are uploaded; each row becomes one AIO, column headers become field names, and values are "
    "stored in up to 50 named element slots.",
    bold_phrases=["[Key.Value]", "50 named element slots"])

add_table(doc,
    headers=["Property", "Specification"],
    rows=[
        ("Storage table",       "aio_data"),
        ("Element capacity",    "50 fields  (element_1 … element_50)"),
        ("Primary key",         "UUID  (aio_id)"),
        ("Identifier",          "aio_name  (indexed)"),
        ("Source formats",      "CSV rows; PDF-extracted structured data"),
        ("Canonical format",    "[FieldName.Value][Field2.Value2]…"),
        ("Tenant isolation",    "Row-level security via app.tenant_id"),
    ],
    col_widths=[2.2, 3.8])

h2(doc, "3.2  Hyper-Semantic Layer (HSL)")

body(doc,
    "HSLs are the platform's relational index. Each HSL record aggregates all AIOs that share "
    "a common element value — creating a precomputed neighborhood for any key/value pair in the "
    "corpus. HSL element slots can additionally hold MRO back-links in [MRO.<uuid>] format, "
    "embedding prior retrieval episodes directly into the semantic index. This is the mechanism "
    "by which the MRO reuse pipeline operates.",
    bold_phrases=["[MRO.<uuid>]", "MRO reuse pipeline"])

add_table(doc,
    headers=["Property", "Specification"],
    rows=[
        ("Storage table",    "hsl_data"),
        ("Element capacity", "100 pointer slots  (hsl_element_1 … hsl_element_100)"),
        ("Primary key",      "UUID  (hsl_id)"),
        ("Content",          "AIO references, MRO back-links [MRO.<uuid>], relationship metadata"),
        ("Generation",       "Manual (Semantic Processor) or auto-rebuild from full AIO corpus"),
        ("Lookup cost",      "Substring search O(HSL_count × 100) — no index recompute"),
    ],
    col_widths=[2.2, 3.8])

h2(doc, "3.3  Memory Result Objects (MROs)")

body(doc,
    "MROs are episodic records of complete AI retrieval sessions. Every ChatAIO response can be "
    "saved as an MRO capturing the full episode: the original natural-language query, the parsed "
    "cue set, the HSL traversal path, the AIO evidence bundle, and the synthesized answer. Once "
    "saved, an MRO can be embedded in HSL element slots, causing future queries that traverse "
    "those HSLs to surface the prior answer as Tier-1 context.",
    bold_phrases=["full episode", "Tier-1 context"])

add_table(doc,
    headers=["Column", "Type", "Description"],
    rows=[
        ("mro_id",             "UUID PK",     "Primary identifier"),
        ("mro_key",            "text indexed","Format: mro-{timestamp}-{random}"),
        ("query_text",         "text",        "Original natural-language query"),
        ("search_terms",       "JSONB",       "Parsed field/value cue set"),
        ("result_text",        "text",        "LLM synthesized answer"),
        ("context_bundle",     "text",        "Full serialized AIO evidence used"),
        ("seed_hsls",          "text",        "Pipe-separated HSL names/IDs traversed"),
        ("matched_aios_count", "integer",     "Traversal cost metric"),
        ("confidence",         "text",        "'derived' or numeric 0..1"),
        ("tenant_id",          "text",        "Tenant isolation key"),
    ],
    col_widths=[1.8, 1.2, 3.0])

body(doc,
    "MRO Ranking Formula — Prior MROs are ranked for injection by compound relevance score:",
    bold_phrases=["MRO Ranking Formula"])
code_block(doc,
    "  score = Jaccard( K_query , K_mro ) × exp( −λ · age_days ) × confidence_m\n\n"
    "  where  K_query = cue set of the current query\n"
    "         K_mro   = cue set captured at MRO creation\n"
    "         λ       = freshness decay constant\n"
    "         confidence_m = MRO confidence rating")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — AI SEARCH ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "4.  AI Search and Retrieval Architecture")

body(doc,
    "The platform implements three distinct search modes, each offering a different cost/precision "
    "profile. All three modes share the same conversational thread in the ChatAIO interface.")

h2(doc, "4.1  Mode 1 — Broad Chat")

body(doc,
    "The full AIO corpus (up to 300 AIO lines) and HSL graph (up to 10 HSL blocks) are fetched "
    "from the database and injected as the Claude system prompt. Suitable for exploratory queries "
    "across the full dataset. Token cost is high and fixed — the entire corpus is sent on every "
    "request regardless of query specificity.")

add_table(doc,
    headers=["Attribute", "Value"],
    rows=[
        ("Endpoint",     "POST /v1/op/chat"),
        ("Context size", "Up to 300 AIO lines + 10 HSL blocks"),
        ("Cost profile", "Fixed full-corpus scan — high token cost"),
        ("Use case",     "Exploratory / broad coverage questions"),
    ], col_widths=[2.0, 4.0])

h2(doc, "4.2  Mode 2 — AIO Search (Four-Phase Search Algebra)")

body(doc,
    "A focused retrieval pipeline that uses the HSL as a pre-filter before calling the LLM, "
    "dramatically reducing context size and token cost while improving answer precision.")

for step, text in [
    ("Phase 1 — Parse",
     "Claude extracts structured search terms (field names + values) from the natural-language "
     "query. Known field names from the Information Elements directory are provided as a hint list. "
     "Output: a structured cue set {(field, value)}."),
    ("Phase 2 — Match HSLs",
     "Every extracted cue is searched across all 100-element HSL records by substring match. "
     "Matching HSLs and their embedded MRO back-links are collected."),
    ("Phase 3 — Gather AIOs",
     "AIOs referenced in matched HSLs are fetched. If no HSLs match, the system falls back to "
     "direct ILIKE search across all 50 AIO element columns."),
    ("Phase 4 — Synthesize",
     "Matched AIOs + ranked MRO priors (from HSL back-links) are assembled and sent to Claude "
     "with the original query. The LLM works from a focused, relevant context rather than the "
     "full corpus."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(step + ":  ")
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5); r.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.name = "Calibri"

doc.add_paragraph()

h2(doc, "4.3  Mode 3 — Substrate Chat (Paper III Pipeline)")

body(doc,
    "The most sophisticated mode, implementing the full Paper III precomputed substrates concept. "
    "Context assembly occurs entirely on the client — no raw database dump is sent to the LLM. "
    "Retrieval is deterministic and reproducible: the same query produces the same cue set, "
    "the same traversal, and the same context bundle.",
    bold_phrases=["entirely on the client", "deterministic and reproducible"])

for step, text in [
    ("Step 1 — Cue Extraction",
     "The query is parsed into a structured cue set K = {(field, value)} using pure lexical "
     "matching against the field and value vocabularies. No LLM is called at this step — the "
     "extraction is fully deterministic."),
    ("Step 2 — HSL Traversal",
     "Set-intersection algebra: N(K) = ⋂_{k∈K} H(k), where H(k) is the AIO neighborhood "
     "for cue k. The intersection identifies AIOs that satisfy all query cues simultaneously."),
    ("Step 3 — MRO Prior Ranking",
     "Prior MROs with overlapping cue sets are scored by the Jaccard × freshness × confidence "
     "formula. The top 5 ranked priors are selected for Tier-1 injection."),
    ("Step 4 — Bundle Assembly",
     "Tier 1: Ranked MRO priors (previously validated answers). "
     "Tier 2: AIOs from HSL traversal (directly matching evidence). "
     "Tier 3: Seed AIOs from the HSL roots (broader contextual support)."),
    ("Step 5 — Claude Inference",
     "The serialized context bundle + conversation history are sent to /v1/op/substrate-chat. "
     "The backend calls Claude with the pre-built bundle only — no database queries occur at "
     "inference time."),
    ("Step 6 — MRO Capture",
     "The session is automatically saved as a new MRO and linked back to the traversed HSLs, "
     "enriching the graph for future queries — the compounding knowledge flywheel."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(step + ":  ")
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5); r.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.name = "Calibri"

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — APPLICATION FEATURES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "5.  Application Features and User Interface")

h2(doc, "5.1  Core Workflows")

for wf, desc in [
    ("Workflow 1 — Data Ingestion",
     "CSV files are uploaded via drag-and-drop or file picker. The AIO Converter previews the "
     "table, optionally maps column headers to AIO field names using Claude field-mapping assist, "
     "and batch-saves each row as an AIO record. PDF files are processed through Claude vision "
     "to produce structured CSV first, then converted to AIOs."),
    ("Workflow 2 — HSL Construction",
     "The Semantic Processor provides element-level navigation: the operator selects any element "
     "value from the corpus, the system finds all AIOs containing that value, and the operator "
     "creates an HSL record linking them. Auto-rebuild generates the full HSL graph from the "
     "entire AIO corpus in a single operation."),
    ("Workflow 3 — AI Query (ChatAIO)",
     "The full-screen ChatAIO interface offers three search modes with a shared conversation "
     "thread. Responses render with full markdown including formatted tables. The operator can "
     "save any response as an MRO, link the MRO to HSLs, export the session as PDF, or download "
     "the transcript as markdown."),
    ("Workflow 4 — MRO Reuse",
     "Once MROs are linked into HSLs, future AIO Search and Substrate queries automatically "
     "surface them as Tier-1 prior context. The system compounds knowledge over time — each "
     "answered query that is saved improves future answers on related topics."),
]:
    h2(doc, wf)
    body(doc, desc)

h2(doc, "5.2  System Administration")

body(doc, "A 10-tab administration panel provides complete platform management:")

add_table(doc,
    headers=["Tab", "Function"],
    rows=[
        ("Users",         "Create, edit, deactivate users; assign System Admin or General User roles"),
        ("Roles",         "Define and remove role categories"),
        ("AIO Data",      "Browse, create, edit, delete all AIO records; inspect element values"),
        ("HSL Data",      "Browse HSLs; dark-blue element badges; Edit / Structure / Delete per record"),
        ("API Key",       "Set and rotate the Anthropic API key (masked display on read)"),
        ("Saved CSVs",    "Library of uploaded CSV source files"),
        ("Saved AIOs",    "Persist converted AIO sets as named collections"),
        ("Saved Prompts", "User-managed library of reusable query prompt templates"),
        ("Info Elements", "Directory of field names with AIO occurrence counts; corpus rebuild"),
        ("Architecture",  "Live system topology diagram and technical reference"),
    ],
    col_widths=[1.5, 4.5])

h2(doc, "5.3  Analytics and Telemetry")

body(doc,
    "The chat_stats table captures detailed session telemetry for every ChatAIO query, enabling "
    "cost analysis, usage pattern review, and cross-mode performance comparison:",
    bold_phrases=["chat_stats"])

for item in [
    "Search mode (Broad / AIO Search / Substrate) and query text",
    "Wall-clock elapsed milliseconds per query",
    "Input and output token counts (cost attribution)",
    "Matched HSL count, matched AIO count, MRO prior count",
    "Cue count and neighborhood size (Substrate mode)",
    "Whether an MRO was saved in the session",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "6.  Technology Stack")

add_table(doc,
    headers=["Layer", "Component", "Technology"],
    rows=[
        ("Frontend",  "Framework",     "Next.js 16.1.6 (React 19, App Router)"),
        ("Frontend",  "Styling",       "Tailwind CSS 4.1 + PostCSS"),
        ("Frontend",  "UI Components", "Radix UI / Shadcn — 40+ accessible primitives"),
        ("Frontend",  "HTTP Client",   "Native Fetch API — typed wrappers in lib/api-client.ts"),
        ("Frontend",  "Package Mgr",   "pnpm (standalone output for deployment)"),
        ("Backend",   "Framework",     "FastAPI (Python 3.10+, Pydantic v2 validation)"),
        ("Backend",   "DB Driver",     "psycopg — raw SQL, no ORM"),
        ("Backend",   "Auth",          "bcrypt password hashing; role-based access control"),
        ("Backend",   "AI SDK",        "Anthropic Python SDK (Claude Sonnet 4.6)"),
        ("Database",  "DBMS",          "PostgreSQL 15 with RLS, JSONB, pg_trgm, btree_gin"),
        ("Database",  "Migrations",    "14 SQL migration files — auto-applied at startup"),
        ("Desktop",   "Wrapper",       "Electron 28  (DMG / NSIS / AppImage)"),
        ("Desktop",   "Bundled",       "Python 3.12 runtime + PostgreSQL binaries"),
        ("AI Model",  "Primary LLM",   "Claude Sonnet 4.6 (claude-sonnet-4-6)"),
        ("AI Model",  "Doc Extraction","Claude Vision — PDF → structured CSV"),
    ],
    col_widths=[1.2, 1.6, 3.2])

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DEPLOYMENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "7.  Deployment Architecture")

h2(doc, "7.1  Cloud Production (Railway)")

add_table(doc,
    headers=["Service", "Description", "Health Check"],
    rows=[
        ("frontend",        "Next.js standalone server — auto-deploy from GitHub main",  "GET /"),
        ("infophysics-api", "FastAPI backend — rootDirectory: infophysics_impl_grade",   "GET /v1/health"),
        ("postgres",        "Managed PostgreSQL 15 — network-isolated",                  "—"),
    ],
    col_widths=[1.6, 3.0, 1.4])

h2(doc, "7.2  Desktop Application (Electron)")

body(doc,
    "The platform packages as a native desktop application using Electron 28. The distributable "
    "bundles the entire stack for fully air-gapped operation — no internet connection required "
    "after installation:")

for item in [
    "Standalone Python 3.12 runtime with all FastAPI dependencies pre-installed",
    "FastAPI backend started as a managed subprocess by the Electron main process",
    "Next.js standalone frontend server on port 3100",
    "PostgreSQL binaries (initdb, pg_ctl, postgres) with self-managed data directory",
    "Schema migrations applied automatically on first launch",
]:
    bullet(doc, item)

add_table(doc,
    headers=["Platform", "Format", "Architecture"],
    rows=[
        ("macOS",   ".dmg",      "arm64 (Apple Silicon) + x64 (Intel)"),
        ("Windows", ".exe",      "NSIS installer — x64"),
        ("Linux",   ".AppImage", "x64"),
    ],
    col_widths=[1.4, 1.4, 3.2])

h2(doc, "7.3  Docker (Self-hosted Enterprise)")

body(doc,
    "docker compose up --build launches all three services (PostgreSQL + FastAPI + Next.js) "
    "in isolated containers, suitable for self-hosted enterprise or air-gapped data-center "
    "deployment.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SECURITY AND MULTI-TENANCY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "8.  Security and Multi-Tenancy")

add_table(doc,
    headers=["Area", "Implementation"],
    rows=[
        ("Authentication",       "Email + password; bcrypt hashing; role-based access (System Admin / General User)"),
        ("Tenant isolation",     "X-Tenant-Id request header; PostgreSQL RLS sets app.tenant_id per transaction"),
        ("Row-level security",   "All 14 core tables protected by RLS policies (migration 003)"),
        ("API key storage",      "Anthropic key stored in system_settings (JSONB); masked on GET response"),
        ("Input validation",     "Pydantic v2 models enforce schema on all FastAPI endpoints"),
        ("Transport security",   "HTTPS via Railway TLS (cloud); localhost-only for desktop"),
        ("CORS",                 "Configurable allowed origins (default: localhost:3000, localhost:3003)"),
    ],
    col_widths=[2.0, 4.0])

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — INTELLECTUAL PROPERTY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "9.  Intellectual Property — Proprietary Concepts")

body(doc,
    "The following concepts are defined, documented, and implemented in this platform. Each "
    "represents a novel contribution with no direct commercial equivalent. All are proprietary "
    "trade secrets of InformationPhysics.ai, LLC.",
    bold_phrases=["proprietary trade secrets of InformationPhysics.ai, LLC"])

ip_items = [
    ("Information Physics Standard Model",
     "A framework for representing, linking, and retrieving structured information through "
     "three deterministic layers (AIO → HSL → MRO), independent of embedding-based similarity. "
     "Documented in Papers I, II, and III."),
    ("Hyper-Semantic Layer (HSL)",
     "A precomputed pointer structure indexing AIO neighborhoods by element value. Enables "
     "O(1) graph traversal at query time with no index recomputation cost as the corpus grows."),
    ("Memory Result Objects (MROs)",
     "Episodic capture of complete retrieval episodes (query + cues + traversal + context + answer) "
     "as first-class database objects linked back into the HSL graph as [MRO.<uuid>] pointers."),
    ("Paper III — Precomputed Substrates",
     "A retrieval paradigm in which LLM context bundles are assembled through deterministic "
     "set-algebra (no embeddings) and delivered as structured substrates. Enables reproducible, "
     "auditable, bounded-cost inference across any corpus size."),
    ("MRO Ranking Formula",
     "score = Jaccard(K_m, K) × exp(−λ · age_days) × confidence_m — a compound relevance "
     "function for surfacing prior retrieval episodes ranked by topical similarity, freshness, "
     "and recorded confidence."),
    ("Tiered Context Assembly",
     "A three-tier context construction strategy (Tier 1: MRO priors, Tier 2: traversal-matched "
     "AIOs, Tier 3: seed AIOs) that prioritizes previously validated answers over raw data, "
     "improving answer quality as the MRO graph grows."),
    ("Self-Improving Knowledge Graph",
     "Each MRO capture and HSL linkage makes the graph denser and more responsive to future "
     "related queries. Knowledge compounds with use — without retraining or re-embedding any model."),
    ("Deterministic Cue Extraction",
     "Natural-language queries are parsed into structured {(field, value)} cue sets using "
     "lexical matching against corpus vocabulary — no LLM required at the extraction step, "
     "yielding reproducible, auditable input to the traversal algebra."),
]

for i, (title, desc) in enumerate(ip_items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"  IP-{i+1:02d}  {title}")
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5); r.font.name = "Calibri"
    body(doc, f"         {desc}")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — KEY METRICS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "10.  Key Metrics and System Constraints")

add_table(doc,
    headers=["Parameter", "Value"],
    rows=[
        ("AIO element capacity",        "50 fields per record"),
        ("HSL element capacity",        "100 pointer slots per record"),
        ("MRO prior injection limit",   "Top 5 ranked priors per query"),
        ("Broad search context limit",  "300 AIO lines + 10 HSL blocks"),
        ("AIO Search fallback",         "Full ILIKE across 50 element columns"),
        ("PDF upload limit",            "20 MB"),
        ("Response max tokens",         "2,048 (chat); 8,192 (PDF extract)"),
        ("Cue parse max tokens",        "500 (search term extraction)"),
        ("Database migrations",         "14 SQL files — auto-applied at startup"),
        ("API endpoints (FastAPI)",     "~45 endpoints"),
        ("Frontend proxy routes",       "~25 Next.js route handlers"),
        ("Admin panel tabs",            "10 management tabs"),
        ("Default tenant",             "tenantA  (RLS enforced)"),
        ("Platform version",            "V3.5  —  April 2026"),
    ],
    col_widths=[3.0, 3.0])

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "11.  Summary Assessment")

body(doc,
    "The Information Physics Demo System is a coherent, production-deployed implementation of "
    "a genuinely novel AI retrieval architecture. Its differentiation from the commercial market "
    "rests on five pillars:")

for item in [
    ("Determinism",
     "Same query produces the same cue set, the same traversal, and the same context bundle. "
     "No probabilistic similarity drift across runs or corpus changes."),
    ("Auditability",
     "Every retrieval episode is recorded in full. The complete path from query to answer — "
     "cues, HSL traversal, matched AIOs, MRO priors, synthesized answer — is traceable and "
     "reproducible."),
    ("Cost Bounding",
     "Context size is a function of HSL neighborhood size, not total corpus size. Cost does "
     "not grow linearly with data volume as it does in embedding-based systems."),
    ("Compounding Value",
     "MRO linkage means the system improves measurably with use, without retraining or "
     "re-embedding any model. Each answered query that is saved enriches future answers on "
     "related topics."),
    ("Deployment Flexibility",
     "The same codebase runs as a Railway cloud service, a self-hosted Docker stack, or a "
     "fully air-gapped native desktop application — serving cloud, enterprise, and secure "
     "facility use cases from a single implementation."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    r = p.add_run(item[0] + ":  ")
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10.5); r.font.name = "Calibri"
    r2 = p.add_run(item[1])
    r2.font.size = Pt(10.5); r2.font.name = "Calibri"

doc.add_paragraph()
body(doc,
    "The intellectual property is original, documented in proprietary reference papers, and "
    "has no direct commercial equivalent. The technology is implemented end-to-end in a modern "
    "production stack — Next.js 16, FastAPI, PostgreSQL 15, Electron 28, Claude Sonnet 4.6 — "
    "under active version-controlled development with CI/CD on Railway.")

doc.add_paragraph()
rule(doc, "CC0000")

# Final confidentiality footer
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot_p.add_run(
    "TRADE SECRET AND PROPRIETARY INFORMATION OF INFORMATIONPHYSICS.AI, LLC\n"
    "© 2026 InformationPhysics.ai, LLC. All rights reserved.\n"
    "Unauthorized disclosure or reproduction is prohibited under applicable trade secret law."
)
r.bold = True
r.font.size  = Pt(8.5)
r.font.name  = "Calibri"
r.font.color.rgb = RED_TS

# ── Save ────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")
